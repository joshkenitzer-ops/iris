import base64
import io
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.tools.docx_checks import (
    check_contact_not_in_header_footer,
    check_em_dash_in_docx,
    check_font_compliance,
    check_full_ats_scan,
    check_hidden_text_in_docx,
    check_illegal_characters,
    check_no_graphics_or_special_heading_chars,
    check_no_tables_or_columns,
    check_physical_formatting,
    check_plain_text_roundtrip,
)
from app.tools.docx_render import generate_resume_docx


def _b64(docx_bytes: bytes) -> str:
    return base64.b64encode(docx_bytes).decode("ascii")


def _clean_docx_b64() -> str:
    return _b64(generate_resume_docx([("Experience", "Led the platform migration.")]))


# Smallest valid PNG (1x1, transparent), hardcoded rather than generated
# via Pillow, so this test suite has no dependency beyond python-docx
# itself, which is already required for everything else in this file.
_MINIMAL_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB\x60\x82"
)


class TestHiddenText(unittest.TestCase):
    def test_clean_document_passes(self) -> None:
        result = check_hidden_text_in_docx(_clean_docx_b64())
        self.assertTrue(result.passed)

    def test_hidden_run_fails(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "visible text")])))
        p = doc.add_paragraph("secret keyword stuffing")
        p.runs[0].font.hidden = True
        buf = io.BytesIO()
        doc.save(buf)
        result = check_hidden_text_in_docx(_b64(buf.getvalue()))
        self.assertFalse(result.passed)
        self.assertEqual(result.findings[0]["severity"], "Critical")

    def test_white_on_white_fails(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "visible text")])))
        p = doc.add_paragraph("invisible keywords")
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        buf = io.BytesIO()
        doc.save(buf)
        result = check_hidden_text_in_docx(_b64(buf.getvalue()))
        self.assertFalse(result.passed)

    def test_near_zero_font_size_fails(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "visible text")])))
        p = doc.add_paragraph("tiny keywords")
        p.runs[0].font.size = Pt(0.5)
        buf = io.BytesIO()
        doc.save(buf)
        result = check_hidden_text_in_docx(_b64(buf.getvalue()))
        self.assertFalse(result.passed)


class TestNoTablesOrColumns(unittest.TestCase):
    def test_clean_document_passes(self) -> None:
        result = check_no_tables_or_columns(_clean_docx_b64())
        self.assertTrue(result.passed)

    def test_table_fails(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "x")])))
        doc.add_table(rows=1, cols=2)
        buf = io.BytesIO()
        doc.save(buf)
        result = check_no_tables_or_columns(_b64(buf.getvalue()))
        self.assertFalse(result.passed)
        self.assertEqual(result.findings[0]["severity"], "Critical")


class TestGraphicsAndHeadingChars(unittest.TestCase):
    def test_clean_document_passes(self) -> None:
        result = check_no_graphics_or_special_heading_chars(_clean_docx_b64())
        self.assertTrue(result.passed)

    def test_embedded_image_fails(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "x")])))
        doc.add_picture(io.BytesIO(_MINIMAL_PNG))
        buf = io.BytesIO()
        doc.save(buf)
        result = check_no_graphics_or_special_heading_chars(_b64(buf.getvalue()))
        self.assertFalse(result.passed)

    def test_special_character_in_heading_fails(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "x")])))
        doc.add_heading("\u2605 Skills", level=1)
        buf = io.BytesIO()
        doc.save(buf)
        result = check_no_graphics_or_special_heading_chars(_b64(buf.getvalue()))
        self.assertFalse(result.passed)


class TestContactNotInHeaderFooter(unittest.TestCase):
    def test_clean_document_passes(self) -> None:
        result = check_contact_not_in_header_footer(_clean_docx_b64())
        self.assertTrue(result.passed)

    def test_email_in_header_fails(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "x")])))
        doc.sections[0].header.paragraphs[0].text = "jane@example.com"
        buf = io.BytesIO()
        doc.save(buf)
        result = check_contact_not_in_header_footer(_b64(buf.getvalue()))
        self.assertFalse(result.passed)

    def test_phone_in_footer_fails(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "x")])))
        doc.sections[0].footer.paragraphs[0].text = "555-0100-1234"
        buf = io.BytesIO()
        doc.save(buf)
        result = check_contact_not_in_header_footer(_b64(buf.getvalue()))
        self.assertFalse(result.passed)


class TestFontCompliance(unittest.TestCase):
    def test_clean_document_passes(self) -> None:
        result = check_font_compliance(_clean_docx_b64())
        self.assertTrue(result.passed)

    def test_disallowed_font_fails(self) -> None:
        docx_bytes = generate_resume_docx([("Experience", "x")], font_name="Comic Sans MS")
        result = check_font_compliance(_b64(docx_bytes))
        self.assertFalse(result.passed)


class TestPhysicalFormatting(unittest.TestCase):
    def test_clean_document_passes(self) -> None:
        result = check_physical_formatting(_clean_docx_b64())
        self.assertTrue(result.passed)

    def test_wrong_margin_fails(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "x")])))
        doc.sections[0].left_margin = Inches(1.5)
        buf = io.BytesIO()
        doc.save(buf)
        result = check_physical_formatting(_b64(buf.getvalue()))
        self.assertFalse(result.passed)

    def test_center_aligned_body_fails(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "x")])))
        p = doc.add_paragraph("centered text")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        buf = io.BytesIO()
        doc.save(buf)
        result = check_physical_formatting(_b64(buf.getvalue()))
        self.assertFalse(result.passed)


class TestPlainTextRoundtrip(unittest.TestCase):
    def test_matching_content_passes(self) -> None:
        docx_bytes = generate_resume_docx([("Experience", "Led the platform migration.")])
        result = check_plain_text_roundtrip(_b64(docx_bytes), ["Experience", "Led the platform migration."])
        self.assertTrue(result.passed)

    def test_dropped_content_fails(self) -> None:
        docx_bytes = generate_resume_docx([("Experience", "Led the platform migration.")])
        result = check_plain_text_roundtrip(
            _b64(docx_bytes), ["Experience", "Led the platform migration.", "Education"]
        )
        self.assertFalse(result.passed)
        self.assertEqual(result.findings[0]["severity"], "Critical")


class TestEmDashInDocx(unittest.TestCase):
    def test_clean_document_passes(self) -> None:
        result = check_em_dash_in_docx(_clean_docx_b64())
        self.assertTrue(result.passed)

    def test_em_dash_in_body_fails(self) -> None:
        docx_bytes = generate_resume_docx([("Experience", "Led migration\u2014end to end.")])
        result = check_em_dash_in_docx(_b64(docx_bytes))
        self.assertFalse(result.passed)


class TestIllegalCharacters(unittest.TestCase):
    def test_clean_document_passes(self) -> None:
        result = check_illegal_characters(_clean_docx_b64())
        self.assertTrue(result.passed)

    def test_curly_quotes_flagged(self) -> None:
        docx_bytes = generate_resume_docx([("Experience", "Led the \u201cplatform\u201d migration.")])
        result = check_illegal_characters(_b64(docx_bytes))
        self.assertFalse(result.passed)


class TestFullAtsScan(unittest.TestCase):
    def test_clean_document_passes(self) -> None:
        result = check_full_ats_scan(_clean_docx_b64())
        self.assertTrue(result.passed)

    def test_table_surfaces_in_composite_findings(self) -> None:
        doc = Document(io.BytesIO(generate_resume_docx([("Experience", "x")])))
        doc.add_table(rows=1, cols=2)
        buf = io.BytesIO()
        doc.save(buf)
        result = check_full_ats_scan(_b64(buf.getvalue()))
        self.assertFalse(result.passed)
        self.assertTrue(any("table" in f["issue"].lower() for f in result.findings))


if __name__ == "__main__":
    unittest.main()
