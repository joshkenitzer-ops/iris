import base64
import io
import unittest

from docx import Document
from docx.shared import Inches, Pt

from app.tools.page_estimate import estimate_page_count, estimate_remaining_page_space


def _docx_base64(bullet_count: int, font_size: int = 11, line_spacing=None) -> str:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    name = doc.add_paragraph("Jane Doe")
    name.runs[0].font.size = Pt(14)
    doc.add_paragraph("Senior Backend Engineer")
    doc.add_paragraph("jane@example.com | 555-123-4567 | Buffalo, NY")
    doc.add_paragraph("SUMMARY")
    doc.add_paragraph(
        "Backend engineer with a decade of experience building distributed "
        "systems at scale, with deep expertise in cloud infrastructure."
    )
    doc.add_paragraph("EXPERIENCE")
    for i in range(bullet_count):
        p = doc.add_paragraph(
            f"Led a cross-functional initiative number {i} that reduced "
            "infrastructure costs by a meaningful double-digit percentage "
            "while improving reliability across every downstream service."
        )
        for run in p.runs:
            run.font.size = Pt(font_size)
        if line_spacing is not None:
            p.paragraph_format.line_spacing = line_spacing

    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode()


class TestEstimatePageCount(unittest.TestCase):
    def test_short_resume_estimates_one_page(self) -> None:
        result = estimate_page_count(_docx_base64(bullet_count=3))
        self.assertTrue(result.passed)
        self.assertEqual(result.data["estimated_page_count"], 1)

    def test_page_count_grows_monotonically_with_content(self) -> None:
        counts = [estimate_page_count(_docx_base64(bullet_count=n)).data["estimated_page_count"] for n in (3, 15, 40)]
        self.assertEqual(counts, sorted(counts))
        self.assertLess(counts[0], counts[-1])

    def test_excessively_long_resume_flags_low_severity(self) -> None:
        result = estimate_page_count(_docx_base64(bullet_count=40))
        self.assertFalse(result.passed)
        self.assertEqual(result.findings[0]["severity"], "Low")
        self.assertIn("heuristic", result.findings[0]["issue"].lower())

    def test_within_target_range_passes(self) -> None:
        result = estimate_page_count(_docx_base64(bullet_count=8))
        self.assertTrue(result.passed)
        self.assertEqual(result.data["target_range"], [1, 2])

    def test_result_includes_target_range(self) -> None:
        result = estimate_page_count(_docx_base64(bullet_count=3))
        self.assertEqual(result.data["target_range"], [1, 2])

    def test_larger_font_size_increases_estimated_pages(self) -> None:
        small = estimate_page_count(_docx_base64(bullet_count=20, font_size=10)).data["estimated_page_count"]
        large = estimate_page_count(_docx_base64(bullet_count=20, font_size=14)).data["estimated_page_count"]
        self.assertGreaterEqual(large, small)

    def test_explicit_line_spacing_multiplier_increases_estimated_pages(self) -> None:
        single = estimate_page_count(_docx_base64(bullet_count=20, line_spacing=1.0)).data["estimated_page_count"]
        double = estimate_page_count(_docx_base64(bullet_count=20, line_spacing=2.0)).data["estimated_page_count"]
        self.assertGreaterEqual(double, single)


class TestEstimateRemainingPageSpace(unittest.TestCase):
    def test_never_gates_even_when_far_over_target(self) -> None:
        result = estimate_remaining_page_space(_docx_base64(bullet_count=60))
        self.assertTrue(result.passed)  # informational only, per module docstring

    def test_short_resume_has_more_remaining_space_than_a_full_one(self) -> None:
        short = estimate_remaining_page_space(_docx_base64(bullet_count=3))
        fuller = estimate_remaining_page_space(_docx_base64(bullet_count=14))
        self.assertGreaterEqual(
            short.data["estimated_page_fraction_remaining"],
            fuller.data["estimated_page_fraction_remaining"],
        )

    def test_remaining_fraction_is_between_zero_and_one(self) -> None:
        for n in (1, 10, 30):
            result = estimate_remaining_page_space(_docx_base64(bullet_count=n))
            fraction = result.data["estimated_page_fraction_remaining"]
            self.assertGreaterEqual(fraction, 0.0)
            self.assertLessEqual(fraction, 1.0)

    def test_result_reports_the_same_page_count_as_estimate_page_count(self) -> None:
        b64 = _docx_base64(bullet_count=15)
        page_count_result = estimate_page_count(b64).data["estimated_page_count"]
        remaining_result = estimate_remaining_page_space(b64).data["estimated_page_count"]
        self.assertEqual(page_count_result, remaining_result)


if __name__ == "__main__":
    unittest.main()
