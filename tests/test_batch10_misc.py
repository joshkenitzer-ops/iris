import base64
import importlib.util
import io
import unittest

from docx import Document

from app.session import Session
from app.tools.intake import ingest_document, validate_structured_intake_form

_PYPDF_AVAILABLE = importlib.util.find_spec("pypdf") is not None
_SKIP_NO_PYPDF = "pypdf is not installed; the PDF path in T-0.1 is optional by design (lazy-imported), and so is testing it."


def _docx_bytes(paragraphs=None, table_rows=None) -> bytes:
    doc = Document()
    for text in paragraphs or []:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row_values in enumerate(table_rows):
            for c, value in enumerate(row_values):
                table.rows[r].cells[c].text = value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _text_pdf_bytes(text: str) -> bytes:
    """Builds a minimal real-text single-page PDF using only pypdf's own
    object model (no reportlab dependency, since pypdf is already the
    hard dependency T-0.1's PDF path relies on)."""
    from pypdf import PdfWriter
    from pypdf.generic import DictionaryObject, NameObject, StreamObject

    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    page = writer.pages[0]

    content = f"BT /F1 24 Tf 100 200 Td ({text}) Tj ET".encode()
    stream = StreamObject()
    stream.set_data(content)
    stream_ref = writer._add_object(stream)

    font = DictionaryObject()
    font[NameObject("/Type")] = NameObject("/Font")
    font[NameObject("/Subtype")] = NameObject("/Type1")
    font[NameObject("/BaseFont")] = NameObject("/Helvetica")
    font_ref = writer._add_object(font)

    font_dict = DictionaryObject()
    font_dict[NameObject("/F1")] = font_ref
    resources = DictionaryObject()
    resources[NameObject("/Font")] = font_dict

    page[NameObject("/Contents")] = stream_ref
    page[NameObject("/Resources")] = resources

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _blank_pdf_bytes() -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _session_with_attachment(filename: str, file_type: str, raw: bytes):
    """Mirrors what POST /sessions/{id}/attachments does: store the raw
    bytes server-side and hand the model back only a short id - never
    the bytes themselves as a tool argument (see
    app.session.Attachment)."""
    session = Session(session_id="s", user_id="u")
    attachment = session.add_attachment(
        filename=filename, file_type=file_type, data=raw
    )
    return session, attachment.id


class TestIngestDocumentDocx(unittest.TestCase):
    def test_extracts_paragraphs_and_table_cells(self) -> None:
        raw = _docx_bytes(
            paragraphs=["John Smith", "Software Engineer with 10 years experience."],
            table_rows=[["Skill", "Python"]],
        )
        session, attachment_id = _session_with_attachment("resume.docx", "docx", raw)
        result = ingest_document(attachment_id, session=session)
        self.assertTrue(result.passed)
        self.assertIn("John Smith", result.data["extracted_text"])
        self.assertIn("Python", result.data["extracted_text"])
        self.assertEqual(result.data["paragraph_count"], 4)

    def test_empty_docx_is_a_high_finding_not_a_silent_empty_string(self) -> None:
        session, attachment_id = _session_with_attachment("empty.docx", "docx", _docx_bytes(paragraphs=[]))
        result = ingest_document(attachment_id, session=session)
        self.assertFalse(result.passed)
        self.assertEqual(result.findings[0]["severity"], "High")
        self.assertEqual(result.data["extracted_text"], "")

    def test_whitespace_only_paragraphs_count_as_empty(self) -> None:
        raw = _docx_bytes(paragraphs=["   ", "\t"])
        session, attachment_id = _session_with_attachment("resume.docx", "docx", raw)
        result = ingest_document(attachment_id, session=session)
        self.assertFalse(result.passed)


@unittest.skipUnless(_PYPDF_AVAILABLE, _SKIP_NO_PYPDF)
class TestIngestDocumentPdf(unittest.TestCase):
    def test_extracts_real_text(self) -> None:
        raw = _text_pdf_bytes("Jane Doe Senior Engineer")
        session, attachment_id = _session_with_attachment("resume.pdf", "pdf", raw)
        result = ingest_document(attachment_id, session=session)
        self.assertTrue(result.passed)
        self.assertIn("Jane Doe Senior Engineer", result.data["extracted_text"])
        self.assertEqual(result.data["page_count"], 1)

    def test_blank_pdf_with_no_text_layer_is_a_high_finding(self) -> None:
        session, attachment_id = _session_with_attachment("resume.pdf", "pdf", _blank_pdf_bytes())
        result = ingest_document(attachment_id, session=session)
        self.assertFalse(result.passed)
        self.assertEqual(result.findings[0]["severity"], "High")
        self.assertIn("scan", result.findings[0]["issue"].lower())


@unittest.skipIf(_PYPDF_AVAILABLE, "this test exercises the path taken specifically when pypdf is absent")
class TestIngestDocumentPdfWithoutPypdf(unittest.TestCase):
    def test_missing_pypdf_is_a_clean_critical_finding_not_a_crash(self) -> None:
        """Pins the actual design intent: pypdf is lazy-imported inside
        _ingest_pdf precisely so an environment without it still starts
        up fine and every docx-only path still works; a PDF ingest
        attempt should decline cleanly, not raise ImportError up
        through the tool call."""
        session, attachment_id = _session_with_attachment("resume.pdf", "pdf", b"irrelevant")
        result = ingest_document(attachment_id, session=session)
        self.assertFalse(result.passed)
        self.assertEqual(result.findings[0]["severity"], "Critical")
        self.assertIn("pypdf", result.findings[0]["issue"].lower())


class TestIngestDocumentInputHandling(unittest.TestCase):
    def test_unknown_attachment_id_is_a_critical_finding_not_a_crash(self) -> None:
        session = Session(session_id="s", user_id="u")
        result = ingest_document("not-a-real-attachment-id", session=session)
        self.assertFalse(result.passed)
        self.assertEqual(result.findings[0]["severity"], "Critical")
        self.assertIn("No attachment", result.findings[0]["issue"])

    def test_corrupt_docx_bytes_are_a_critical_finding_not_a_crash(self) -> None:
        """Replaces the corrupted-stored-base64 test that guarded this
        same property until attachments stopped being stored base64
        (2026-07-27). The encoding changed; a truncated or mislabeled
        upload must still fail cleanly with something the user can act
        on, rather than throwing an opaque tool error."""
        session = Session(session_id="s", user_id="u")
        attachment = session.add_attachment(filename="x.docx", file_type="docx", data=b"not-a-real-docx!!!")
        result = ingest_document(attachment.id, session=session)
        self.assertFalse(result.passed)
        self.assertEqual(result.findings[0]["severity"], "Critical")
        self.assertIn("x.docx", result.findings[0]["issue"])

    def test_corrupt_pdf_bytes_are_a_critical_finding_not_a_crash(self) -> None:
        session = Session(session_id="s", user_id="u")
        attachment = session.add_attachment(filename="x.pdf", file_type="pdf", data=b"not-a-real-pdf!!!")
        result = ingest_document(attachment.id, session=session)
        self.assertFalse(result.passed)
        self.assertEqual(result.findings[0]["severity"], "Critical")

    def test_unsupported_file_type_is_rejected(self) -> None:
        raw = _docx_bytes(paragraphs=["text"])
        session, attachment_id = _session_with_attachment("resume.txt", "txt", raw)
        result = ingest_document(attachment_id, session=session)
        self.assertFalse(result.passed)
        self.assertIn("Unsupported file_type", result.findings[0]["issue"])

    def test_attachments_are_isolated_per_session(self) -> None:
        raw = _docx_bytes(paragraphs=["text"])
        session_a, attachment_id = _session_with_attachment("resume.docx", "docx", raw)
        session_b = Session(session_id="other", user_id="u2")
        result = ingest_document(attachment_id, session=session_b)
        self.assertFalse(result.passed)
        self.assertIn("No attachment", result.findings[0]["issue"])


class TestSessionAttachments(unittest.TestCase):
    """Session.add_attachment/get_attachment: the storage half of the
    mechanism ingest_document (T-0.1) reads from. The HTTP upload
    endpoint itself (POST /sessions/{id}/attachments) is tested in
    tests/test_main_integration.py, since it needs a live FastAPI
    TestClient; this covers the storage logic those requests land on."""

    def test_added_attachment_is_retrievable_by_id(self) -> None:
        session = Session(session_id="s", user_id="u")
        attachment = session.add_attachment("resume.docx", "docx", "aGVsbG8=")
        retrieved = session.get_attachment(attachment.id)
        self.assertIs(retrieved, attachment)
        self.assertEqual(retrieved.filename, "resume.docx")
        self.assertEqual(retrieved.file_type, "docx")

    def test_unknown_id_returns_none_not_a_crash(self) -> None:
        session = Session(session_id="s", user_id="u")
        self.assertIsNone(session.get_attachment("does-not-exist"))

    def test_each_attachment_gets_a_distinct_id(self) -> None:
        session = Session(session_id="s", user_id="u")
        a = session.add_attachment("a.docx", "docx", b"hello")
        b = session.add_attachment("b.docx", "docx", b"hello")
        self.assertNotEqual(a.id, b.id)

    def test_quota_evicts_the_oldest_attachment_first(self) -> None:
        from app.config import MAX_ATTACHMENTS_PER_SESSION

        session = Session(session_id="s", user_id="u")
        ids = [session.add_attachment(f"f{i}.docx", "docx", b"hello").id for i in range(MAX_ATTACHMENTS_PER_SESSION + 5)]
        self.assertLessEqual(len(session.attachments), MAX_ATTACHMENTS_PER_SESSION)
        self.assertNotIn(ids[0], session.attachments)  # oldest, evicted
        self.assertIn(ids[-1], session.attachments)  # newest, kept

    def test_attachments_do_not_leak_between_sessions(self) -> None:
        session_a = Session(session_id="a", user_id="u1")
        session_b = Session(session_id="b", user_id="u2")
        attachment = session_a.add_attachment("resume.docx", "docx", b"hello")
        self.assertIsNone(session_b.get_attachment(attachment.id))


class TestSessionAttachmentByteBudget(unittest.TestCase):
    """The count cap alone let a session hold far more memory than the
    instance can spare: ten files just under the per-file limit satisfied
    it while adding up to ~140 MB against a 512 MB box. These pin the
    byte budget that closes that hole (2026-07-27)."""

    def test_byte_budget_evicts_oldest_until_the_newcomer_fits(self) -> None:
        from app.config import MAX_ATTACHMENT_BYTES_PER_SESSION

        session = Session(session_id="s", user_id="u")
        chunk = MAX_ATTACHMENT_BYTES_PER_SESSION // 4
        ids = [session.add_attachment(f"f{i}.docx", "docx", b"x" * chunk).id for i in range(6)]

        self.assertLessEqual(session.attachment_bytes(), MAX_ATTACHMENT_BYTES_PER_SESSION)
        self.assertNotIn(ids[0], session.attachments)  # oldest, evicted
        self.assertIn(ids[-1], session.attachments)  # newest, always kept

    def test_budget_holds_under_the_count_cap(self) -> None:
        """Well under MAX_ATTACHMENTS_PER_SESSION by count, still over
        budget by bytes: the case a count-only cap missed entirely."""
        from app.config import MAX_ATTACHMENT_BYTES_PER_SESSION, MAX_ATTACHMENTS_PER_SESSION

        session = Session(session_id="s", user_id="u")
        for i in range(4):
            session.add_attachment(f"f{i}.docx", "docx", b"x" * (MAX_ATTACHMENT_BYTES_PER_SESSION // 3))

        self.assertLess(len(session.attachments), MAX_ATTACHMENTS_PER_SESSION)
        self.assertLessEqual(session.attachment_bytes(), MAX_ATTACHMENT_BYTES_PER_SESSION)

    def test_cached_extracted_text_counts_against_the_budget(self) -> None:
        """A text-heavy PDF's extraction can rival the file itself; a
        budget that counted only stored file bytes would undercount it."""
        session = Session(session_id="s", user_id="u")
        attachment = session.add_attachment("f.pdf", "pdf", b"x" * 1000)
        before = session.attachment_bytes()
        attachment.extracted_text = "y" * 5000
        self.assertEqual(session.attachment_bytes(), before + 5000)

    def test_a_single_attachment_is_always_stored(self) -> None:
        """Eviction must not loop forever or drop the only file when one
        upload is large relative to the budget."""
        from app.config import MAX_ATTACHMENT_BYTES_PER_SESSION

        session = Session(session_id="s", user_id="u")
        attachment = session.add_attachment("big.pdf", "pdf", b"x" * (MAX_ATTACHMENT_BYTES_PER_SESSION + 1))
        self.assertIn(attachment.id, session.attachments)
        self.assertEqual(len(session.attachments), 1)


class TestValidateStructuredIntakeForm(unittest.TestCase):
    def test_all_required_fields_present_passes(self) -> None:
        result = validate_structured_intake_form(
            {"name": "Josh", "email": "j@example.com"}, ["name", "email"]
        )
        self.assertTrue(result.passed)

    def test_missing_field_is_critical(self) -> None:
        result = validate_structured_intake_form({"name": "Josh"}, ["name", "email", "target_role"])
        self.assertFalse(result.passed)
        self.assertEqual(result.data["missing_fields"], ["email", "target_role"])
        self.assertTrue(all(f["severity"] == "Critical" for f in result.findings))

    def test_blank_field_counts_as_missing(self) -> None:
        result = validate_structured_intake_form({"name": "  ", "email": "j@example.com"}, ["name", "email"])
        self.assertFalse(result.passed)
        self.assertEqual(result.data["missing_fields"], ["name"])

    def test_extra_unrequired_fields_are_ignored(self) -> None:
        result = validate_structured_intake_form(
            {"name": "Josh", "email": "j@x.com", "hobby": "chess"}, ["name", "email"]
        )
        self.assertTrue(result.passed)

    def test_empty_required_fields_list_always_passes(self) -> None:
        result = validate_structured_intake_form({}, [])
        self.assertTrue(result.passed)


if __name__ == "__main__":
    unittest.main()
