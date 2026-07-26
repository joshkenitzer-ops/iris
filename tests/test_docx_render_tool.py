import io
import unittest

from docx import Document

from app.session import Session
from app.tools.docx_render import render_resume_docx_tool


class TestRenderResumeDocxTool(unittest.TestCase):
    def setUp(self) -> None:
        self.session = Session(session_id="s", user_id="u")

    def test_produces_valid_docx_and_stores_on_session(self) -> None:
        result = render_resume_docx_tool(
            sections=[{"heading": "Experience", "body": "Led the migration."}],
            filename="Test_User_Resume_Acme_Eng_V1.docx",
            session=self.session,
        )
        self.assertTrue(result.passed)
        # Tool now returns file_id/filename, not raw base64
        self.assertIn("file_id", result.data)
        self.assertEqual(result.data["filename"], "Test_User_Resume_Acme_Eng_V1.docx")
        # The actual bytes live in the session
        rendered = self.session.get_rendered_file(result.data["file_id"])
        self.assertIsNotNone(rendered)
        import base64
        docx_bytes = base64.b64decode(rendered.data_base64)
        doc = Document(io.BytesIO(docx_bytes))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        self.assertIn("Experience", texts)
        self.assertIn("Led the migration.", texts)

    def test_custom_font_applied(self) -> None:
        result = render_resume_docx_tool(
            sections=[{"heading": "Experience", "body": "Led the migration."}],
            filename="Test_User_Resume_Acme_Eng_V1.docx",
            session=self.session,
            font_name="Georgia",
        )
        self.assertTrue(result.passed)
        rendered = self.session.get_rendered_file(result.data["file_id"])
        import base64
        docx_bytes = base64.b64decode(rendered.data_base64)
        doc = Document(io.BytesIO(docx_bytes))
        body_para = [p for p in doc.paragraphs if p.text == "Led the migration."][0]
        for run in body_para.runs:
            self.assertEqual(run.font.name, "Georgia")


if __name__ == "__main__":
    unittest.main()
