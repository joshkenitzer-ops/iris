import base64
import unittest

from docx import Document
import io

from app.tools.docx_render import generate_resume_docx


class TestGenerateResumeDocx(unittest.TestCase):
    def test_produces_valid_docx_bytes(self) -> None:
        docx_bytes = generate_resume_docx([("Experience", "Led the platform migration.")])
        doc = Document(io.BytesIO(docx_bytes))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        self.assertIn("Experience", texts)
        self.assertIn("Led the platform migration.", texts)

    def test_margins_are_one_inch(self) -> None:
        from docx.shared import Inches

        docx_bytes = generate_resume_docx([("Experience", "x")])
        doc = Document(io.BytesIO(docx_bytes))
        section = doc.sections[0]
        self.assertEqual(section.left_margin, Inches(1))
        self.assertEqual(section.right_margin, Inches(1))
        self.assertEqual(section.top_margin, Inches(1))
        self.assertEqual(section.bottom_margin, Inches(1))

    def test_font_applied_to_runs(self) -> None:
        docx_bytes = generate_resume_docx([("Experience", "Led the migration.")], font_name="Georgia")
        doc = Document(io.BytesIO(docx_bytes))
        body_paragraph = [p for p in doc.paragraphs if p.text == "Led the migration."][0]
        for run in body_paragraph.runs:
            self.assertEqual(run.font.name, "Georgia")

    def test_multiple_sections_preserve_order(self) -> None:
        docx_bytes = generate_resume_docx([("Experience", "First."), ("Education", "Second.")])
        doc = Document(io.BytesIO(docx_bytes))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        self.assertEqual(texts, ["Experience", "First.", "Education", "Second."])

    def test_produces_base64_encodable_bytes(self) -> None:
        docx_bytes = generate_resume_docx([("Experience", "x")])
        encoded = base64.b64encode(docx_bytes).decode("ascii")
        decoded = base64.b64decode(encoded)
        self.assertEqual(decoded, docx_bytes)


if __name__ == "__main__":
    unittest.main()
