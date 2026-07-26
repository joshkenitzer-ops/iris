"""
Phase 0 intake tools.

T-0.2 is HYBRID: the scoring itself, character ratios, role-block
counts, date-parse failure rates, is fully deterministic and lives
here. What stays judgment, per the spec, is whether content that
scores as low-confidence is still usable; this module does not decide
that, it only produces the score and the specific signals that
tripped so a human or a later JUDGMENT pass has something concrete to
look at.

T-0.1 (below) does the actual document parsing and feeds
score_extraction_confidence with real numbers instead of test
fixtures; it does not compute confidence itself, only text and page/
paragraph counts.
"""

from __future__ import annotations

import base64
import io
import logging
from typing import List

from app.config import EXTRACTION_CONFIDENCE, MAX_INGEST_TEXT_CHARS
from app.enforcement import EnforcementKind, ToolResult, registry, tool
from app.session import Attachment, Session
from app.untrusted_text import wrap_untrusted

logger = logging.getLogger("iris.intake")


@tool(
    id="T-0.1",
    name="ingest_document",
    description=(
        "Extracts raw text from a previously uploaded docx or PDF, "
        "referenced by attachment_id (returned by POST "
        "/sessions/{session_id}/attachments when the file was "
        "uploaded). Parsing is solved; model tokens spent "
        "transcribing a document by hand, or even just typing out its "
        "base64 bytes as a tool argument, are waste, which is exactly "
        "why this takes a short reference instead of file content. "
        "This is deterministic extraction only, no interpretation of "
        "what the text means. Feeds text and a paragraph/page count "
        "forward into score_extraction_confidence (T-0.2), which does "
        "the actual confidence scoring. A PDF with no extractable "
        "text layer, most likely a scan with no OCR pass, surfaces as "
        "a High finding rather than silently returning an empty "
        "string, since that case needs routing to manual review "
        "(T-0.3), not a Phase 1 pass built on nothing."
    ),
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {"attachment_id": {"type": "string"}},
        "required": ["attachment_id"],
    },
    needs_session=True,
)
def ingest_document(attachment_id: str, session: Session) -> ToolResult:
    attachment = session.get_attachment(attachment_id)
    if attachment is None:
        return ToolResult(
            passed=False,
            findings=[
                {
                    "severity": "Critical",
                    "issue": f"No attachment with id '{attachment_id}' on this session.",
                    "fix": "Upload the file via POST /sessions/{session_id}/attachments and use the returned attachment_id.",
                }
            ],
        )

    try:
        raw = base64.b64decode(attachment.file_base64)
    except Exception:
        return ToolResult(
            passed=False,
            findings=[{"severity": "Critical", "issue": "Stored attachment is not valid base64.", "fix": "Upload the file again."}],
        )

    if attachment.file_type == "docx":
        return _ingest_docx(raw, attachment)
    if attachment.file_type == "pdf":
        return _ingest_pdf(raw, attachment)
    return ToolResult(
        passed=False,
        findings=[{"severity": "Critical", "issue": f"Unsupported file_type '{attachment.file_type}'.", "fix": "Use 'docx' or 'pdf'."}],
    )


def _ingest_docx(raw: bytes, attachment: Attachment) -> ToolResult:
    from docx import Document  # already a hard dependency, via app.tools.docx_render/docx_checks

    document = Document(io.BytesIO(raw))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    table_count = len(document.tables)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    text = "\n".join(paragraphs)

    logger.info(
        "ingest_document (docx): %d paragraphs (incl. table cells), %d tables, %d chars total extracted",
        len(paragraphs),
        table_count,
        len(text),
    )

    if not text.strip():
        logger.warning("ingest_document (docx): extraction produced no usable text (empty after strip)")
        return ToolResult(
            passed=False,
            findings=[
                {
                    "severity": "High",
                    "issue": "No extractable text found in the docx.",
                    "fix": "Route to manual review (T-0.3); the file may be empty or image-only.",
                }
            ],
            data={"extracted_text": "", "paragraph_count": 0},
        )
    # Cached raw (unwrapped, same truncation as the model-facing copy) so
    # run_batch_checks can resolve `text` server-side from attachment_id
    # instead of the model having to paste tens of thousands of characters
    # back into a tool call just to name a document it already extracted.
    attachment.extracted_text = text[:MAX_INGEST_TEXT_CHARS]
    return ToolResult(
        passed=True,
        data={
            "extracted_text": wrap_untrusted(text, "uploaded .docx"),
            "paragraph_count": len(paragraphs),
            "raw_char_count": len(text),
        },
    )


def _ingest_pdf(raw: bytes, attachment: Attachment) -> ToolResult:
    try:
        from pypdf import PdfReader  # lazy import: PDF support is optional, must not break docx-only startup
    except ImportError:
        return ToolResult(
            passed=False,
            findings=[
                {
                    "severity": "Critical",
                    "issue": "pypdf is not installed in this environment.",
                    "fix": "Install pypdf, or route this PDF to manual review (T-0.3) instead.",
                }
            ],
        )

    reader = PdfReader(io.BytesIO(raw))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages_text)

    logger.info(
        "ingest_document (pdf): %d pages, %d chars total extracted",
        len(reader.pages),
        len(text),
    )

    if not text.strip():
        logger.warning("ingest_document (pdf): extraction produced no usable text (empty after strip) — likely scan/no OCR")
        return ToolResult(
            passed=False,
            findings=[
                {
                    "severity": "High",
                    "issue": "No extractable text layer in the PDF; likely a scan with no OCR.",
                    "fix": "Route to manual review (T-0.3) rather than treating this as a complete extraction.",
                }
            ],
            data={"extracted_text": "", "page_count": len(reader.pages)},
        )
    attachment.extracted_text = text[:MAX_INGEST_TEXT_CHARS]
    return ToolResult(
        passed=True,
        data={
            "extracted_text": wrap_untrusted(text, "uploaded .pdf"),
            "page_count": len(reader.pages),
            "raw_char_count": len(text),
        },
    )


@tool(
    id="T-0.2",
    name="score_extraction_confidence",
    description=(
        "Scores extraction confidence from four deterministic signals: "
        "OCR/text-layer confidence, the ratio of replacement or control "
        "characters, the count of role blocks with a parseable date "
        "range, and the date-parse failure ratio. Returns pass/fail per "
        "signal against the configured thresholds (app.config."
        "EXTRACTION_CONFIDENCE) plus which signals tripped. Does not "
        "decide whether low-confidence content is still usable, that "
        "judgment happens after this returns."
    ),
    kind=EnforcementKind.HYBRID,
    input_schema={
        "type": "object",
        "properties": {
            "ocr_confidence": {"type": "number", "description": "0.0-1.0. Use 1.0 if there is a native text layer."},
            "replacement_char_ratio": {"type": "number", "description": "Fraction of characters that are replacement/control chars."},
            "role_blocks_with_dates": {"type": "integer"},
            "date_parse_failure_ratio": {"type": "number", "description": "Fraction of date-like strings that failed to parse."},
        },
        "required": [
            "ocr_confidence",
            "replacement_char_ratio",
            "role_blocks_with_dates",
            "date_parse_failure_ratio",
        ],
    },
)
def score_extraction_confidence(
    ocr_confidence: float,
    replacement_char_ratio: float,
    role_blocks_with_dates: int,
    date_parse_failure_ratio: float,
) -> ToolResult:
    # These four values are supplied BY THE MODEL as tool-call arguments,
    # not measured directly by this function or by ingest_document — this
    # tool only applies thresholds to whatever numbers it's handed. Logging
    # them raw is essential: a "100% date parse failure" in the user-facing
    # message could reflect a real signal computed from actual extracted
    # text, or the model estimating/asserting a plausible-looking number
    # without having done the underlying parsing. This log line is what
    # lets that distinction be checked after the fact.
    logger.info(
        "score_extraction_confidence inputs: ocr_confidence=%.3f replacement_char_ratio=%.4f "
        "role_blocks_with_dates=%d date_parse_failure_ratio=%.3f",
        ocr_confidence,
        replacement_char_ratio,
        role_blocks_with_dates,
        date_parse_failure_ratio,
    )

    thresholds = EXTRACTION_CONFIDENCE
    tripped = []

    if ocr_confidence < thresholds["ocr_min_confidence"]:
        tripped.append(f"ocr_confidence {ocr_confidence:.2f} below {thresholds['ocr_min_confidence']}")
    if replacement_char_ratio > thresholds["max_replacement_char_ratio"]:
        tripped.append(f"replacement_char_ratio {replacement_char_ratio:.3f} above {thresholds['max_replacement_char_ratio']}")
    if role_blocks_with_dates < thresholds["min_role_blocks_with_dates"]:
        tripped.append(f"role_blocks_with_dates {role_blocks_with_dates} below {thresholds['min_role_blocks_with_dates']}")
    if date_parse_failure_ratio > thresholds["max_date_parse_failure_ratio"]:
        tripped.append(f"date_parse_failure_ratio {date_parse_failure_ratio:.2f} above {thresholds['max_date_parse_failure_ratio']}")

    passed = len(tripped) == 0
    logger.info("score_extraction_confidence result: passed=%s tripped=%s", passed, tripped)
    findings = (
        []
        if passed
        else [
            {
                "severity": "High",
                "issue": f"Extraction confidence signal(s) tripped: {'; '.join(tripped)}.",
                "fix": "Route to manual review (T-0.3) rather than treating this extraction as complete.",
            }
        ]
    )
    return ToolResult(passed=passed, findings=findings, data={"tripped_signals": tripped})


@tool(
    id="T-0.3",
    name="route_low_confidence_to_manual_review",
    description=(
        "Gate on score_extraction_confidence's result. Any tripped "
        "signal routes to manual review; this never decides the "
        "extraction is 'probably fine anyway'. Fail-closed by design."
    ),
    kind=EnforcementKind.GATE,
    input_schema={
        "type": "object",
        "properties": {
            "extraction_passed": {"type": "boolean"},
            "tripped_signals": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["extraction_passed"],
    },
    blocking=True,
)
def route_low_confidence_to_manual_review(extraction_passed: bool, tripped_signals=None) -> ToolResult:
    if extraction_passed:
        return ToolResult(passed=True)
    return ToolResult(
        passed=False,
        findings=[
            {
                "severity": "High",
                "issue": "Extraction confidence too low to proceed automatically.",
                "fix": "Route to manual review or another entry path before Phase 1.",
            }
        ],
        data={"tripped_signals": tripped_signals or []},
    )


# ---------------------------------------------------------------------------
# T-9.14: get current date for master resume filename generation
# ---------------------------------------------------------------------------


@tool(
    id="T-9.14",
    name="get_todays_date",
    description=(
        "Returns today's date in YYYY-MM-DD format for use in master "
        "resume filenames. The spec requires the master resume to use "
        "the current date at generation time, not any date embedded in "
        "the user's uploaded source document. Always call this tool "
        "when constructing the master resume filename rather than "
        "deriving a date from the uploaded file."
    ),
    kind=EnforcementKind.TOOL,
    input_schema={"type": "object", "properties": {}},
)
def get_todays_date() -> ToolResult:
    from datetime import datetime as _dt

    return ToolResult(
        passed=True,
        data={"date": _dt.now().strftime("%Y-%m-%d")},
    )


# ---------------------------------------------------------------------------
# T-0.6: careerInventory schema validation
# ---------------------------------------------------------------------------

# Locked relative order, spec Phase 2. HEADLINE and SUMMARY are
# Iris-generated: never inventory-required, always output-required.
_SECTION_ORDER = [
    "NAME",
    "HEADLINE",
    "CONTACT",
    "SUMMARY",
    "SKILLS",
    "EXPERIENCE",
    "EDUCATION",
    "PROJECTS",
    "PUBLICATIONS",
]
_INVENTORY_REQUIRED = {"NAME", "CONTACT", "SKILLS", "EXPERIENCE"}
_GENERATED_SECTIONS = {"HEADLINE", "SUMMARY"}  # never inventory-required by definition


@tool(
    id="T-0.6",
    name="check_career_inventory_schema",
    description=(
        "Validates a careerInventory payload against the locked schema: "
        "no fixed section count, but NAME/CONTACT/SKILLS/EXPERIENCE are "
        "inventory-required, HEADLINE and SUMMARY are Iris-generated "
        "(never required of the user), and everything else is optional. "
        "Flags missing required sections and any section name outside "
        "the known set."
    ),
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {
            "sections": {
                "type": "object",
                "description": "Map of section name to its content string. Absent or empty means not supplied.",
                "additionalProperties": {"type": "string"},
            }
        },
        "required": ["sections"],
    },
)
def check_career_inventory_schema(sections: dict) -> ToolResult:
    findings = []

    unknown = [name for name in sections if name not in _SECTION_ORDER]
    for name in unknown:
        findings.append(
            {
                "severity": "Medium",
                "issue": f"Unrecognized section '{name}'.",
                "fix": f"Use one of: {', '.join(_SECTION_ORDER)}.",
            }
        )

    for required in _INVENTORY_REQUIRED:
        if not sections.get(required, "").strip():
            findings.append(
                {
                    "severity": "Critical",
                    "issue": f"Inventory-required section '{required}' is missing or empty.",
                    "fix": "Phase 2 cannot complete without this section.",
                }
            )

    for generated in _GENERATED_SECTIONS:
        if sections.get(generated):
            findings.append(
                {
                    "severity": "Low",
                    "issue": f"'{generated}' was supplied directly, but it is Iris-generated per spec Phase 2.",
                    "fix": "Confirm this wasn't meant to seed generation rather than be treated as final.",
                }
            )

    return ToolResult(passed=len(findings) == 0, findings=findings, data={"unknown_sections": unknown})


@tool(
    id="T-0.7",
    name="validate_structured_intake_form",
    description=(
        "Validates a structured intake form submission: every name in "
        "required_fields must be present in fields and non-blank. "
        "'Form. No model.' is the entire spec note for T-0.7, so this "
        "is deliberately generic, plain field-presence validation, "
        "distinct from check_career_inventory_schema (T-0.6), which "
        "governs resume section content specifically, not whatever "
        "other structured intake questions Phase 0 asks."
    ),
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {
            "fields": {
                "type": "object",
                "additionalProperties": {"type": "string"},
            },
            "required_fields": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["fields", "required_fields"],
    },
)
def validate_structured_intake_form(fields: dict, required_fields: List[str]) -> ToolResult:
    missing = [name for name in required_fields if not fields.get(name, "").strip()]
    findings = [
        {
            "severity": "Critical",
            "issue": f"Required intake field '{name}' is missing or blank.",
            "fix": "Prompt the user for this field before proceeding.",
        }
        for name in missing
    ]
    return ToolResult(passed=len(missing) == 0, findings=findings, data={"missing_fields": missing})


# ---------------------------------------------------------------------------
# T-0.8: near-duplicate collapse in bulk sources
# ---------------------------------------------------------------------------

from difflib import SequenceMatcher  # noqa: E402


@tool(
    id="T-0.8",
    name="find_near_duplicate_candidates",
    description=(
        "Nominates near-duplicate claim/bullet pairs from a bulk import "
        "using string similarity. Deliberately over-inclusive: a later "
        "instance can be a genuine revision rather than a copy, so this "
        "only nominates candidates above the similarity threshold. "
        "Choosing the canonical version, or confirming a pair are "
        "actually distinct, is judgment this tool does not attempt."
    ),
    kind=EnforcementKind.HYBRID,
    input_schema={
        "type": "object",
        "properties": {
            "items": {"type": "array", "items": {"type": "string"}},
            "threshold": {"type": "number", "description": "Similarity ratio, 0.0-1.0. Default 0.85."},
        },
        "required": ["items"],
    },
)
def find_near_duplicate_candidates(items: List[str], threshold: float = 0.85) -> ToolResult:
    candidates = []
    for i in range(len(items)):
        for j in range(i + 1, len(items)):
            ratio = SequenceMatcher(None, items[i], items[j]).ratio()
            if ratio >= threshold:
                candidates.append({"index_a": i, "index_b": j, "similarity": round(ratio, 3)})

    findings = [
        {
            "severity": "Low",
            "issue": f"Items {c['index_a']} and {c['index_b']} are {c['similarity']:.0%} similar.",
            "fix": "Confirm whether this is a duplicate or a genuine revision before keeping both.",
        }
        for c in candidates
    ]
    return ToolResult(passed=len(candidates) == 0, findings=findings, data={"candidates": candidates})


# ---------------------------------------------------------------------------
# T-8.5: missing required sections, checked against OUTPUT-required rather
# than inventory-required. Same underlying disposition table as T-0.6;
# PROJECTS is deliberately excluded, its output presence is a relevance
# call (T-6.16), never a required flag.
# ---------------------------------------------------------------------------

_OUTPUT_REQUIRED = {"NAME", "HEADLINE", "CONTACT", "SUMMARY", "SKILLS", "EXPERIENCE"}


@tool(
    id="T-8.5",
    name="check_missing_required_sections",
    description=(
        "Checks a rendered resume's present sections against the "
        "output-required set (NAME, HEADLINE, CONTACT, SUMMARY, "
        "SKILLS, EXPERIENCE). PROJECTS is never required here, its "
        "inclusion is a relevance decision, not a schema rule."
    ),
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {
            "present_sections": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["present_sections"],
    },
)
def check_missing_required_sections(present_sections: List[str]) -> ToolResult:
    present = set(present_sections)
    missing = sorted(_OUTPUT_REQUIRED - present)
    findings = [
        {
            "severity": "Critical",
            "issue": f"Output-required section '{section}' is missing from the rendered document.",
            "fix": "This section must be present before the document can ship.",
        }
        for section in missing
    ]
    return ToolResult(passed=len(missing) == 0, findings=findings, data={"missing_sections": missing})


# ---------------------------------------------------------------------------
# T-9.15: run multiple TOOL checks in a single harness call
# ---------------------------------------------------------------------------


@tool(
    id="T-9.15",
    name="run_batch_checks",
    description=(
        "Runs multiple TOOL-kind checks in a single harness call, "
        "eliminating per-check model round trips. "
        "IMPORTANT: never pass raw docx_base64 bytes or a pasted-back copy "
        "of extracted_text in inputs — pass attachment_id and the harness "
        "resolves both docx bytes and cached extracted text server-side, "
        "keeping large payloads out of model context and out of the "
        "model's own output entirely. "
        "For Phase 1 (Audit): tool_ids=[T-3.1,T-3.3,T-3.4,T-3.5,T-3.6,T-3.7,T-3.8], "
        "inputs={attachment_id: <id>} (requires ingest_document to have run "
        "on that attachment first, so the text is cached). "
        "For Phase 4 (Formatting): tool_ids=[T-4.1,T-4.2,T-4.3,T-4.4,T-4.9], "
        "inputs={attachment_id: <id>}. "
        "For Phase 8 (Final Review): tool_ids=[T-8.5,T-8.6,T-8.7,T-8.12,T-8.14,T-8.15], "
        "inputs={attachment_id: <id>}. "
        "Only TOOL and GATE kind items are accepted."
    ),
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {
            "tool_ids": {
                "type": "array",
                "items": {"type": "string"},
                "description": "List of T-x.y tool IDs to run in one batch.",
            },
            "inputs": {
                "type": "object",
                "description": (
                    "Shared input dict — a JSON object of key/value pairs, "
                    "e.g. {\"attachment_id\": \"...\"} or {\"text\": \"...\"}. "
                    "Never an array. For docx-based checks, pass "
                    "attachment_id instead of docx_base64 — harness "
                    "resolves bytes server-side."
                ),
                "additionalProperties": True,
                "minProperties": 1,
            },
        },
        "required": ["tool_ids", "inputs"],
    },
    needs_session=True,
)
def run_batch_checks(tool_ids: list, inputs: dict, session: "Session") -> ToolResult:
    """Harness-side batch execution with two payload protections:

    1. attachment_id resolution: if inputs has attachment_id, harness
       resolves docx bytes AND cached extracted text (set by
       ingest_document) from session store and substitutes them before
       dispatching — model never handles raw base64, and never has to
       paste extracted_text back into a tool call just to name the
       document it came from. Each per-tool dispatch (dispatch_by_id)
       already filters inputs down to the keys that tool's function
       signature accepts, so resolving both from one attachment_id is
       safe even though only one is relevant per tool.
    2. data stripping: only passed/findings returned, never tool data
       payloads (candidate lists, extracted text, scores)."""
    from app.enforcement import EnforcementKind as EK

    # Anthropic's SDK deserializes a JSON-schema "object" tool argument as a
    # Python dict, so `inputs` should already be one. But `input_schema`
    # only declares `{"type": "object"}` with no property constraints, so a
    # malformed model call (e.g. inputs supplied as a list of one-element
    # entries instead of key/value pairs) reaches this function unchecked.
    # dict(inputs) on that shape raises ValueError deep inside the batch
    # dispatcher, which aborts the entire batch (all tool_ids in it) with an
    # opaque traceback instead of a finding the model can act on. Validate
    # explicitly and fail closed with a clear, model-legible error instead.
    if not isinstance(inputs, dict):
        return ToolResult(
            passed=False,
            findings=[{
                "severity": "Critical",
                "issue": (
                    f"run_batch_checks received 'inputs' as {type(inputs).__name__}, "
                    "not an object/dict of key-value pairs."
                ),
                "fix": (
                    "Pass inputs as a JSON object, e.g. {\"attachment_id\": \"...\"} "
                    "or {\"text\": \"...\"}, not a list or a single value."
                ),
            }],
        )

    # Resolve attachment_id -> docx_base64 and/or cached extracted_text
    # server-side if provided. Both are attempted from the same id since
    # dispatch_by_id filters inputs to whatever each tool's signature
    # actually accepts (T-3.x text checks read `text`, T-4.x/T-8.x docx
    # checks read `docx_base64`) — there is no cross-contamination.
    inputs = dict(inputs)
    if "attachment_id" in inputs:
        attachment_id = inputs.pop("attachment_id")
        attachment = session.get_attachment(attachment_id)
        if attachment is None:
            return ToolResult(
                passed=False,
                findings=[{
                    "severity": "Critical",
                    "issue": f"No attachment '{attachment_id}' on this session.",
                    "fix": "Upload the file first and use the returned attachment_id.",
                }],
            )
        if "docx_base64" not in inputs:
            inputs["docx_base64"] = attachment.file_base64
        if "text" not in inputs and attachment.extracted_text is not None:
            inputs["text"] = attachment.extracted_text

    per_tool = []
    all_passed = True
    all_findings = []

    for tid in tool_ids:
        try:
            spec = registry.get(tid)
            if spec.kind not in (EK.TOOL, EK.GATE):
                all_findings.append({
                    "severity": "High",
                    "issue": f"{tid} is {spec.kind.name} — call it separately.",
                    "fix": "Remove from batch.",
                })
                per_tool.append({"tool_id": tid, "passed": False})
                all_passed = False
                continue
            result = registry.dispatch_by_id(tid, inputs, session=session)
            per_tool.append({"tool_id": tid, "passed": result.passed})
            all_findings.extend(result.findings)
            if not result.passed:
                all_passed = False
        except Exception as exc:  # noqa: BLE001
            all_findings.append({
                "severity": "Critical",
                "issue": f"Tool {tid} raised: {exc}",
                "fix": "Check server logs.",
            })
            per_tool.append({"tool_id": tid, "passed": False})
            all_passed = False

    return ToolResult(
        passed=all_passed,
        findings=all_findings,
        data={"summary": per_tool},
    )
