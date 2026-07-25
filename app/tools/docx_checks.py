"""
docx-level structural checks. Everything here takes base64-encoded
docx bytes as input, since raw binary has no place in a JSON tool
call; a Claude tool call or an HTTP request carries the file this way,
and _load_document decodes it back into a real python-docx Document
before any check runs.

These are all TOOL or GATE kind: mechanical inspection of the file's
XML structure, never judgment about whether the document reads well.
"""

from __future__ import annotations

import base64
import io
import re
from typing import List

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from app.config import ALLOWED_FONTS, BODY_FONT_SIZE_RANGE, NAME_FONT_SIZE_RANGE
from app.enforcement import EnforcementKind, ToolResult, tool
from app.tools.slop import check_em_dash

_EMAIL_RE = re.compile(r"[^@\s]+@[^@\s]+\.[^@\s]+")
_PHONE_RE = re.compile(r"(\+?\d[\d\-.\s]{7,}\d)")
_SMART_PUNCT = {
    "\u2018": "left single quote",
    "\u2019": "right single quote",
    "\u201c": "left double quote",
    "\u201d": "right double quote",
    "\u2026": "ellipsis character",
    "\u00a0": "non-breaking space",
}
_ALLOWED_HEADING_CHARS = set(
    "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789 .,&/'-"
)


def _load_document(docx_base64: str) -> Document:
    raw = base64.b64decode(docx_base64)
    return Document(io.BytesIO(raw))


def _all_runs(doc: Document):
    for paragraph in doc.paragraphs:
        yield from paragraph.runs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield from paragraph.runs


def _extract_plain_text(doc: Document) -> List[str]:
    return [p.text for p in doc.paragraphs if p.text.strip()]


@tool(
    id="T-3.2",
    name="check_hidden_text_in_docx",
    description=(
        "Positive-visibility check per Design Principle 8: scans every "
        "run in the document, including table cells, for the "
        "hidden/vanish property, near-zero font size, or white font "
        "color, any of which render text invisible to a human reader "
        "while remaining machine-readable, a prompt-injection vector "
        "into any downstream parser. Heuristic on color: assumes a "
        "white page background rather than inspecting cell or "
        "paragraph shading."
    ),
    kind=EnforcementKind.GATE,
    input_schema={
        "type": "object",
        "properties": {"docx_base64": {"type": "string"}},
        "required": ["docx_base64"],
    },
    blocking=True,
)
def check_hidden_text_in_docx(docx_base64: str) -> ToolResult:
    doc = _load_document(docx_base64)
    findings = []
    for run in _all_runs(doc):
        if run.font.hidden:
            findings.append({"severity": "Critical", "issue": f"Hidden run found: '{run.text[:40]}'.", "fix": "Remove the hidden run entirely."})
            continue
        if run.font.size is not None and run.font.size.pt < 1:
            findings.append({"severity": "Critical", "issue": f"Near-zero font size on run: '{run.text[:40]}'.", "fix": "Remove or resize to a legible, visible size."})
            continue
        if run.font.color is not None and run.font.color.rgb is not None and str(run.font.color.rgb) == "FFFFFF":
            findings.append({"severity": "Critical", "issue": f"White-on-white text found: '{run.text[:40]}'.", "fix": "Remove or set a visible color."})
    return ToolResult(passed=len(findings) == 0, findings=findings)


@tool(
    id="T-4.1",
    name="check_no_tables_or_columns",
    description="Checks the document contains no tables and no multi-column section layout.",
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {"docx_base64": {"type": "string"}},
        "required": ["docx_base64"],
    },
)
def check_no_tables_or_columns(docx_base64: str) -> ToolResult:
    doc = _load_document(docx_base64)
    findings = []
    if len(doc.tables) > 0:
        findings.append({"severity": "Critical", "issue": f"{len(doc.tables)} table(s) found.", "fix": "Rebuild using plain paragraphs; ATS parsers frequently fail on tables."})
    for section in doc.sections:
        cols = section._sectPr.find(qn("w:cols"))
        num_attr = cols.get(qn("w:num")) if cols is not None else None
        if num_attr and int(num_attr) > 1:
            findings.append({"severity": "Critical", "issue": f"Section uses {num_attr} columns.", "fix": "Use a single-column layout."})
    return ToolResult(passed=len(findings) == 0, findings=findings)


@tool(
    id="T-4.2",
    name="check_no_graphics_or_special_heading_chars",
    description="Checks for embedded images/shapes and for non-standard characters in heading text.",
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {"docx_base64": {"type": "string"}},
        "required": ["docx_base64"],
    },
)
def check_no_graphics_or_special_heading_chars(docx_base64: str) -> ToolResult:
    doc = _load_document(docx_base64)
    findings = []
    if len(doc.inline_shapes) > 0:
        findings.append({"severity": "Critical", "issue": f"{len(doc.inline_shapes)} embedded graphic(s)/shape(s) found.", "fix": "Remove all graphics, icons, and shapes."})
    for paragraph in doc.paragraphs:
        if not paragraph.style.name.lower().startswith("heading"):
            continue
        bad_chars = sorted({c for c in paragraph.text if c not in _ALLOWED_HEADING_CHARS})
        if bad_chars:
            findings.append(
                {
                    "severity": "High",
                    "issue": f"Heading '{paragraph.text}' contains non-standard character(s): {' '.join(bad_chars)}.",
                    "fix": "Use plain text headings; no icons, emoji, or decorative characters.",
                }
            )
    return ToolResult(passed=len(findings) == 0, findings=findings)


@tool(
    id="T-4.3",
    name="check_contact_not_in_header_footer",
    description="Checks that no email or phone pattern appears in any section's header or footer.",
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {"docx_base64": {"type": "string"}},
        "required": ["docx_base64"],
    },
)
def check_contact_not_in_header_footer(docx_base64: str) -> ToolResult:
    doc = _load_document(docx_base64)
    findings = []
    for i, section in enumerate(doc.sections):
        for label, container in (("header", section.header), ("footer", section.footer)):
            text = " ".join(p.text for p in container.paragraphs)
            if _EMAIL_RE.search(text) or _PHONE_RE.search(text):
                findings.append(
                    {
                        "severity": "Critical",
                        "issue": f"Contact information found in section {i} {label}.",
                        "fix": "Move contact information into the document body; ATS parsers frequently ignore headers and footers.",
                    }
                )
    return ToolResult(passed=len(findings) == 0, findings=findings)


@tool(
    id="T-4.4",
    name="check_font_compliance",
    description=(
        "Checks every run uses an allowed font family. Also checks "
        "body text sizes fall in the 10-12pt band and, treating the "
        "document's first paragraph as the name line, that its size "
        "falls in the 11-14pt band."
    ),
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {"docx_base64": {"type": "string"}},
        "required": ["docx_base64"],
    },
)
def check_font_compliance(docx_base64: str) -> ToolResult:
    doc = _load_document(docx_base64)
    findings = []

    for run in _all_runs(doc):
        font_name = run.font.name or doc.styles["Normal"].font.name
        if font_name and font_name not in ALLOWED_FONTS:
            findings.append({"severity": "High", "issue": f"Font '{font_name}' is not in the allowed list.", "fix": f"Use one of: {', '.join(sorted(ALLOWED_FONTS))}."})

    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        name_para = paragraphs[0]
        for run in name_para.runs:
            size = run.font.size or doc.styles["Normal"].font.size
            if size is not None and not (NAME_FONT_SIZE_RANGE[0] <= size.pt <= NAME_FONT_SIZE_RANGE[1]):
                findings.append({"severity": "Medium", "issue": f"Name line font size {size.pt}pt is outside {NAME_FONT_SIZE_RANGE}.", "fix": "Resize the name line into the allowed range."})
        for para in paragraphs[1:]:
            for run in para.runs:
                size = run.font.size or doc.styles["Normal"].font.size
                if size is not None and not (BODY_FONT_SIZE_RANGE[0] <= size.pt <= BODY_FONT_SIZE_RANGE[1]):
                    findings.append({"severity": "Medium", "issue": f"Body text font size {size.pt}pt is outside {BODY_FONT_SIZE_RANGE}.", "fix": "Resize body text into the allowed range."})

    return ToolResult(passed=len(findings) == 0, findings=findings)


@tool(
    id="T-4.9",
    name="check_physical_formatting",
    description="Checks 1-inch margins on every section and left alignment on body paragraphs.",
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {"docx_base64": {"type": "string"}},
        "required": ["docx_base64"],
    },
)
def check_physical_formatting(docx_base64: str) -> ToolResult:
    doc = _load_document(docx_base64)
    findings = []
    one_inch = Inches(1)

    for i, section in enumerate(doc.sections):
        for side, value in (
            ("left", section.left_margin),
            ("right", section.right_margin),
            ("top", section.top_margin),
            ("bottom", section.bottom_margin),
        ):
            if value != one_inch:
                findings.append(
                    {
                        "severity": "Medium",
                        "issue": f"Section {i} {side} margin is {value / 914400:.2f}in, expected 1in.",
                        "fix": "Set all four margins to exactly 1 inch.",
                    }
                )

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        if paragraph.style.name.lower().startswith("heading"):
            continue
        if paragraph.alignment not in (None, WD_ALIGN_PARAGRAPH.LEFT):
            findings.append({"severity": "Low", "issue": f"Paragraph '{paragraph.text[:40]}' is not left-aligned.", "fix": "Set alignment to left."})

    return ToolResult(passed=len(findings) == 0, findings=findings)


@tool(
    id="T-4.10",
    name="check_plain_text_roundtrip",
    description=(
        "Extracts plain text from a generated docx and compares it "
        "against the expected paragraph sequence, in order. Catches "
        "scrambled characters, merged sections, or dropped fields "
        "introduced during generation."
    ),
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {
            "docx_base64": {"type": "string"},
            "expected_texts": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["docx_base64", "expected_texts"],
    },
)
def check_plain_text_roundtrip(docx_base64: str, expected_texts: List[str]) -> ToolResult:
    doc = _load_document(docx_base64)
    actual = _extract_plain_text(doc)
    expected = [t for t in expected_texts if t.strip()]

    if actual == expected:
        return ToolResult(passed=True, data={"extracted": actual})

    findings = [
        {
            "severity": "Critical",
            "issue": f"Extracted text does not match expected content. Got {len(actual)} paragraph(s), expected {len(expected)}.",
            "fix": "Regenerate the document; content was scrambled, merged, or dropped during rendering.",
        }
    ]
    return ToolResult(passed=False, findings=findings, data={"extracted": actual, "expected": expected})


@tool(
    id="T-8.12",
    name="check_em_dash_in_docx",
    description="Extracts all text from a docx and delegates to the em-dash check (T-3.1). Same rule, docx-level entry point for the Team Lead pass.",
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {"docx_base64": {"type": "string"}},
        "required": ["docx_base64"],
    },
)
def check_em_dash_in_docx(docx_base64: str) -> ToolResult:
    doc = _load_document(docx_base64)
    full_text = "\n".join(_extract_plain_text(doc))
    return check_em_dash(full_text)


@tool(
    id="T-8.13",
    name="check_illegal_characters",
    description="Flags smart/curly quotes, the ellipsis character, and non-breaking spaces, common Word autocorrect artifacts that some ATS parsers mis-render.",
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {"docx_base64": {"type": "string"}},
        "required": ["docx_base64"],
    },
)
def check_illegal_characters(docx_base64: str) -> ToolResult:
    doc = _load_document(docx_base64)
    full_text = "\n".join(_extract_plain_text(doc))
    found = sorted({c for c in full_text if c in _SMART_PUNCT})
    findings = [
        {
            "severity": "Low",
            "issue": f"Illegal character found: {_SMART_PUNCT[c]} (U+{ord(c):04X}).",
            "fix": "Replace with the plain ASCII equivalent.",
        }
        for c in found
    ]
    return ToolResult(passed=len(found) == 0, findings=findings, data={"illegal_characters": found})


@tool(
    id="T-8.15",
    name="check_full_ats_scan",
    description="Composite ATS scan: runs the table/column, graphics/heading, header-footer, and font checks together and aggregates every finding.",
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {"docx_base64": {"type": "string"}},
        "required": ["docx_base64"],
    },
)
def check_full_ats_scan(docx_base64: str) -> ToolResult:
    sub_results = [
        check_no_tables_or_columns(docx_base64),
        check_no_graphics_or_special_heading_chars(docx_base64),
        check_contact_not_in_header_footer(docx_base64),
        check_font_compliance(docx_base64),
    ]
    all_findings = [f for r in sub_results for f in r.findings]
    return ToolResult(passed=all(r.passed for r in sub_results), findings=all_findings)
