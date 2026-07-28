"""
T-4.8: docx render.

Applies the mechanical rules the spec states as absolute: 1-inch
margins, an allowed font, no columns, no tables. It does not make
layout or content decisions, those are JUDGMENT (T-4.15, T-4.16) and
happen before this function is called.

Section content arrives as free text per (heading, body) pair (the
tool's input schema, unchanged). Structure within that text is
recovered here rather than carried as separate fields, since the
locked careerInventory order (spec 5.9) already fixes what each
heading means:

- NAME / HEADLINE / CONTACT render as the resume's letterhead block
  (centered), not as visible section headings.
- SUMMARY / SKILLS / EXPERIENCE / PROJECTS bodies are one line per
  bullet. A bold-lead bullet ("Label: rest", spec 3.1) renders with
  the label bold, matching the check_bold_lead_structure (T-2.2)
  convention already enforced earlier in the pipeline. Within
  EXPERIENCE/PROJECTS, a line shaped like "Title | Org | Mon YYYY -
  Mon YYYY" (matching T-4.5's date format) renders as a bold role
  header instead of a bullet, and the line immediately following it
  is treated as the 1-2 sentence role summary (spec 3.1) if it is not
  itself bullet-shaped.
- Any other heading (EDUCATION, PUBLICATIONS, or a cover letter's
  prose body) renders as blank-line-separated paragraphs, no bullets,
  no bold-lead parsing, since prose sentences may contain colons that
  are not bold-lead labels.

Deliberately out of scope here: page count (T-4.11) and remaining
page-space measurement (T-4.12). Neither is obtainable from
python-docx alone, it does not perform layout, a real page count
needs an actual rendering engine (LibreOffice headless, or a
Word-compatible conversion service) sitting between this generator
and the final check. That is an infrastructure decision, not a
missing function, and it is called out rather than faked with a
character-count estimate, which is exactly the shortcut the spec
warns against.
"""

from __future__ import annotations

import re
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.config import BODY_FONT_SIZE_RANGE
from app.session import Session

DEFAULT_FONT = "Calibri"
DEFAULT_BODY_SIZE = Pt(BODY_FONT_SIZE_RANGE[0] + 1)  # 11pt, mid-range of the 10-12pt band

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)
MUTED_COLOR = RGBColor(0x59, 0x59, 0x59)

_HEADER_BLOCK_HEADINGS = {"NAME", "HEADLINE", "CONTACT"}
_BULLET_SECTION_HEADINGS = {"SUMMARY", "SKILLS", "EXPERIENCE", "PROJECTS"}

_ROLE_HEADER_RE = re.compile(
    r"^.+\|.+\|.*[A-Z][a-z]{2} \d{4}\s*-\s*[A-Z][a-z]{2} \d{4}\s*$"
)
_BOLD_LEAD_RE = re.compile(r"^(?P<lead>[A-Z][^:]{1,60}?):\s+(?P<rest>.+)$")
_MAX_LEAD_WORDS = 8


def _set_run_font(run, font_name: str, size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_header_line(doc, text: str, font_name: str, *, size, bold: bool, color) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _set_run_font(run, font_name, size=size, bold=bold, color=color)


def _add_contact_line(doc, contact_text: str, font_name: str) -> None:
    fields = [f.strip() for f in contact_text.split("|") if f.strip()]
    _add_header_line(doc, "  |  ".join(fields), font_name, size=Pt(9), bold=False, color=MUTED_COLOR)


def _add_section_heading(doc, text: str, font_name: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    _set_run_font(run, font_name, size=Pt(13), bold=True, color=ACCENT_COLOR)


def _add_bold_lead_bullet(doc, line: str, font_name: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    match = _BOLD_LEAD_RE.match(line)
    if match and len(match.group("lead").split()) <= _MAX_LEAD_WORDS:
        lead_run = para.add_run(f"{match.group('lead')}: ")
        _set_run_font(lead_run, font_name, size=DEFAULT_BODY_SIZE, bold=True)
        rest_run = para.add_run(match.group("rest"))
        _set_run_font(rest_run, font_name, size=DEFAULT_BODY_SIZE, bold=False)
    else:
        run = para.add_run(line)
        _set_run_font(run, font_name, size=DEFAULT_BODY_SIZE, bold=False)


def _add_role_header(doc, line: str, font_name: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(line)
    _set_run_font(run, font_name, size=DEFAULT_BODY_SIZE, bold=True)


def _add_role_summary(doc, line: str, font_name: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(line)
    _set_run_font(run, font_name, size=DEFAULT_BODY_SIZE, italic=True, color=MUTED_COLOR)


def _add_plain_paragraph(doc, text: str, font_name: str) -> None:
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        _set_run_font(run, font_name, size=DEFAULT_BODY_SIZE, bold=False)


def _render_bulleted_body(doc, heading: str, body: str, font_name: str) -> None:
    lines = [line.strip() for line in body.splitlines()]
    allow_role_headers = heading.strip().upper() in {"EXPERIENCE", "PROJECTS"}
    previous_was_role_header = False
    for line in lines:
        if not line:
            previous_was_role_header = False
            continue
        if allow_role_headers and _ROLE_HEADER_RE.match(line):
            _add_role_header(doc, line, font_name)
            previous_was_role_header = True
            continue
        if previous_was_role_header and not _BOLD_LEAD_RE.match(line):
            _add_role_summary(doc, line, font_name)
            previous_was_role_header = False
            continue
        _add_bold_lead_bullet(doc, line, font_name)
        previous_was_role_header = False


def _render_prose_body(doc, body: str, font_name: str) -> None:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", body) if p.strip()]
    if not paragraphs:
        paragraphs = [body]
    for paragraph_text in paragraphs:
        collapsed = " ".join(line.strip() for line in paragraph_text.splitlines() if line.strip())
        _add_plain_paragraph(doc, collapsed, font_name)


def generate_resume_docx(sections: List[Tuple[str, str]], font_name: str = DEFAULT_FONT) -> bytes:
    """sections is an ordered list of (heading, body_text) pairs, in
    the locked section order the caller is responsible for supplying
    correctly; this function renders what it's given; ordering is
    T-2.8's job, not this one's."""
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = DEFAULT_BODY_SIZE

    for heading, body in sections:
        key = heading.strip().upper()
        if key == "NAME":
            _add_header_line(doc, body, font_name, size=Pt(18), bold=True, color=ACCENT_COLOR)
        elif key == "HEADLINE":
            _add_header_line(doc, body, font_name, size=DEFAULT_BODY_SIZE, bold=True, color=ACCENT_COLOR)
        elif key == "CONTACT":
            _add_contact_line(doc, body, font_name)
        elif key in _BULLET_SECTION_HEADINGS:
            _add_section_heading(doc, heading, font_name)
            _render_bulleted_body(doc, heading, body, font_name)
        else:
            _add_section_heading(doc, heading, font_name)
            _render_prose_body(doc, body, font_name)

    buffer = _save_to_bytes(doc)
    return buffer


def _save_to_bytes(doc: Document) -> bytes:
    import io

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


# ---------------------------------------------------------------------------
# Claude-callable wrapper. generate_resume_docx above is the real
# implementation and stays plain Python so other code (tests, other
# tools) can call it directly; this registers the same function as a
# tool so the model can trigger rendering itself.
# ---------------------------------------------------------------------------

import base64

from app.enforcement import EnforcementKind, ToolResult, tool
from app.gates import GateBlocked, require_gap_not_silently_removed, require_no_open_criticals
from app.tools.formatting import _FOUNDATIONAL_RE


def _is_deliverable(filename: str) -> bool:
    """Whether rendering this filename constitutes DELIVERY, the moment
    spec rule 4.4 puts the Phase 8 gates at.

    The foundational resume is deliberately excluded. Spec Phase 2 is
    explicit on both halves of why: "the foundational resume is the
    source document, not a document to send," and Iris "immediately
    renders the docx" on Foundational Build completion, which happens
    long before Final Review. Gating every render would therefore break
    Phase 2 outright for a real and common case: a Phase 1 Critical that
    the user acknowledged with a stated reason satisfies
    require_phase1_disposition but is still counted by
    open_criticals(), since dispositioned is not dismissed. Those users
    would have been unable to render a foundational resume at all.

    Anything not recognizable as a foundational resume is treated as a
    deliverable, so an unrecognized filename fails closed rather than
    slipping past the gate. check_filename_pattern (T-4.13) is what
    should have rejected it before reaching here."""
    return not _FOUNDATIONAL_RE.match(filename.strip())


@tool(
    id="T-4.8",
    name="render_resume_docx",
    description=(
        "Renders an ordered list of (heading, body) sections into a "
        "docx file: 1-inch margins, an allowed font, no layout "
        "decisions beyond what's given. NAME/HEADLINE/CONTACT render "
        "as a centered letterhead block. Within SUMMARY/SKILLS/"
        "EXPERIENCE/PROJECTS, put one bullet per line using the "
        "bold-lead 'Label: rest' shape (spec 3.1) for the label to "
        "render bold; in EXPERIENCE/PROJECTS a line shaped 'Title | "
        "Org | Mon YYYY - Mon YYYY' renders as a bold role header, and "
        "the line right after it (if not itself bullet-shaped) is the "
        "italic role summary. Other sections render as blank-line-"
        "separated paragraphs. Stores the file server-side and "
        "returns a file_id the user can download from the UI — the "
        "model never sends raw docx bytes to the user directly. Also "
        "requires a filename that already passed check_filename_pattern "
        "(T-4.13). "
        "Does not decide section order or content, those come from "
        "earlier phases."
    ),
    kind=EnforcementKind.TOOL,
    input_schema={
        "type": "object",
        "properties": {
            "sections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {"heading": {"type": "string"}, "body": {"type": "string"}},
                    "required": ["heading", "body"],
                },
            },
            "filename": {"type": "string", "description": "The validated output filename, e.g. Kenitzer_Joshua_Resume_Oracle_SrIDDev_V1.docx"},
            "font_name": {"type": "string"},
        },
        "required": ["sections", "filename"],
    },
    needs_session=True,
)
def render_resume_docx_tool(sections: list, filename: str, session: Session, font_name: str = DEFAULT_FONT) -> ToolResult:
    pairs = [(s["heading"], s["body"]) for s in sections]

    # The delivery gates, enforced where delivery actually happens.
    #
    # T-8.18 and T-7.8 previously ran only in POST /sessions/{id}/deliver,
    # which nothing in the product ever called: static/app.js has no
    # reference to it, so every session sat in STARTING_POINT and both
    # gates were unreachable in production despite passing their tests
    # (the tests POST the route directly). Meanwhile this function
    # stored a file and the harness emitted file_ready, which the browser
    # turned into a download button, so a user could download a resume
    # with open Critical findings. Spec Principle 9, "programmatic
    # verification," was decorative at runtime as a result.
    #
    # Returning passed=False rather than raising is deliberate: a raise
    # is swallowed by the dispatch layer into a generic "tool failed to
    # run" that tells the model nothing actionable. The real enforcement
    # is that no file is stored, so no file_id comes back, so no
    # download button appears; the findings just make the reason
    # legible to the model and therefore to the user.
    if _is_deliverable(filename):
        final_text = "\n".join(f"{heading}\n{body}" for heading, body in pairs)
        try:
            require_no_open_criticals(session)
            require_gap_not_silently_removed(session, final_text)
        except GateBlocked as exc:
            return ToolResult(
                passed=False,
                findings=[
                    {
                        "severity": "Critical",
                        "issue": f"[{exc.gate_id}] {exc.message}",
                        "fix": (
                            "Resolve the finding, then render again. This document was not "
                            "produced and there is nothing for the user to download yet."
                        ),
                    }
                ],
                data={"blocked_by_gate": exc.gate_id},
            )

    docx_bytes = generate_resume_docx(pairs, font_name=font_name)
    b64 = base64.b64encode(docx_bytes).decode("ascii")

    # Length check, at the one point the rendered document actually
    # exists. estimate_page_count (T-4.11) and its companion T-4.12 were
    # written but called from nowhere: not in any documented batch list,
    # not referenced by Tailoring or Final Review. Confirmed live
    # 2026-07-28, a tailored resume came out at 6 pages against the
    # 1-2 page target with nothing in the pipeline to notice. Same shape
    # as the delivery-gate finding a day earlier, a correct check that
    # was simply never invoked.
    #
    # Advisory, not gating, deliberately. T-4.11 reports Low severity
    # because the estimate is a heuristic with a real margin of error
    # (python-docx does no layout); blocking a render on an approximation
    # would be wrong. The model sees the finding and can act on it.
    #
    # Deliverables only. Spec Phase 2 is explicit that the foundational
    # resume "has no length ceiling; comprehensiveness is its purpose,"
    # so measuring it against a tailored target would flag every one.
    length_findings = []
    if _is_deliverable(filename):
        from app.tools.page_estimate import estimate_page_count

        length_findings = estimate_page_count(b64).findings

    rendered = session.add_rendered_file(
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data_base64=b64,
    )
    # passed stays True even when the length finding fires: the file was
    # produced and is downloadable. The finding is information for the
    # model to act on, not a failure of the render.
    return ToolResult(
        passed=True,
        findings=length_findings,
        data={"file_id": rendered.id, "filename": rendered.filename},
    )
